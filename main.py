from fastapi import FastAPI, File, UploadFile
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
from typing import List
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
from num2words import num2words
import tempfile, math, os, shutil

app = FastAPI(title="Invo GEN API")

TEMPLATE_PATH = "template.docx"

class InvoiceItem(BaseModel):
    description: str
    hsn: str
    qty: float
    rate: float
    unit: str = "Nos"

class InvoiceData(BaseModel):
    customer_name: str
    invoice_no: str
    invoice_date: str
    items: List[InvoiceItem]

@app.post("/generate_invoice")
async def generate_invoice(data: InvoiceData):
    try:
        subtotal = sum(item.qty * item.rate for item in data.items)
        sgst = subtotal * 0.09
        cgst = subtotal * 0.09
        total_tax = sgst + cgst
        grand_total = subtotal + total_tax
        round_off = round(grand_total)
        amount_in_words = num2words(round_off, to='cardinal').title() + " Only"

        doc = Document(TEMPLATE_PATH)

        mapping = {
            "customer_name": data.customer_name,
            "invoice_no": data.invoice_no,
            "invoice_date": data.invoice_date,
            "subtotal": f"{subtotal:.2f}",
            "sgst": f"{sgst:.2f}",
            "cgst": f"{cgst:.2f}",
            "total_tax": f"{total_tax:.2f}",
            "grand_total": f"{grand_total:.2f}",
            "round_off": f"{round_off}",
            "amount_in_words": amount_in_words
        }

        for p in doc.paragraphs:
            for key, val in mapping.items():
                if f"{{{{{key}}}}}" in p.text:
                    inline = p.runs
                    for i in range(len(inline)):
                        if f"{{{{{key}}}}}" in inline[i].text:
                            inline[i].text = inline[i].text.replace(f"{{{{{key}}}}}", str(val))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, val in mapping.items():
                            if f"{{{{{key}}}}}" in p.text:
                                inline = p.runs
                                for i in range(len(inline)):
                                    if f"{{{{{key}}}}}" in inline[i].text:
                                        inline[i].text = inline[i].text.replace(f"{{{{{key}}}}}", str(val))

        if doc.tables:
            items_table = doc.tables[0]
            header_row = items_table.rows[0]

            rows_to_delete = []
            for i in range(1, len(items_table.rows)):
                rows_to_delete.append(items_table.rows[i])

            for row in rows_to_delete:
                items_table._element.remove(row._element)

            for idx, item in enumerate(data.items, start=1):
                row = items_table.add_row()
                cells = row.cells

                if len(cells) >= 6:
                    cells[0].text = str(idx)
                    cells[1].text = item.description
                    cells[2].text = item.hsn
                    cells[3].text = f"{item.qty:.2f}"
                    cells[4].text = item.unit
                    cells[5].text = f"{item.rate:.2f}"
                    if len(cells) >= 7:
                        amount = item.qty * item.rate
                        cells[6].text = f"{amount:.2f}"

        tmp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(tmp_dir, f"invoice_{data.invoice_no}.docx")
        pdf_path = docx_path.replace(".docx", ".pdf")

        doc.save(docx_path)
        convert(docx_path, pdf_path)

        return FileResponse(pdf_path, media_type="application/pdf", filename=f"invoice_{data.invoice_no}.pdf")

    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)

@app.get("/view_template")
async def view_template():
    try:
        if not os.path.exists(TEMPLATE_PATH):
            return JSONResponse(content={"error": "Template file not found"}, status_code=404)
        return FileResponse(TEMPLATE_PATH, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="template.docx")
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)

@app.get("/template_info")
async def template_info():
    try:
        if not os.path.exists(TEMPLATE_PATH):
            return JSONResponse(content={"error": "Template file not found"}, status_code=404)

        doc = Document(TEMPLATE_PATH)
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        absolute_path = os.path.abspath(TEMPLATE_PATH)

        return {
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "template_path": absolute_path
        }
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)

@app.post("/upload_template")
async def upload_template(file: UploadFile = File(...)):
    try:
        if not file.filename.endswith(".docx"):
            return JSONResponse(content={"error": "Only .docx files are allowed"}, status_code=400)

        with open(TEMPLATE_PATH, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        return {"message": "Template uploaded successfully", "filename": file.filename}
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)

@app.get("/")
def index():
    return {"message": "Invo Gen API running with uv"}
